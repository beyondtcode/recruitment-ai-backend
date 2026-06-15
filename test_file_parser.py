"""Unit tests for CV text and hyperlink extraction."""



from __future__ import annotations



from io import BytesIO



from docx import Document

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from pypdf import PdfWriter

from pypdf.generic import (

    ArrayObject,

    DictionaryObject,

    NameObject,

    NumberObject,

    TextStringObject,

)



from utils.file_parser import (

    DOCX_HEADER_END,

    DOCX_HEADER_START,

    _extract_docx,

    _extract_pdf,

    _format_inline_hyperlink,

    _inject_pdf_hyperlinks_into_text,

    _paragraph_text_with_hyperlinks,

    extract_text_from_file,

)





def _add_external_hyperlink(paragraph, text: str, url: str) -> None:

    part = paragraph.part

    r_id = part.relate_to(

        url,

        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",

        is_external=True,

    )

    hyperlink = OxmlElement("w:hyperlink")

    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")

    text_elem = OxmlElement("w:t")

    text_elem.text = text

    run.append(text_elem)

    hyperlink.append(run)

    paragraph._element.append(hyperlink)





def test_format_inline_hyperlink_bracket_notation():

    assert (

        _format_inline_hyperlink("LinkedIn", "https://www.linkedin.com/in/jane-doe")

        == "LinkedIn [https://www.linkedin.com/in/jane-doe]"

    )

    assert (

        _format_inline_hyperlink("", "https://www.linkedin.com/in/john")

        == "LinkedIn [https://www.linkedin.com/in/john]"

    )





def test_docx_inline_hyperlink_bracket_format():

    doc = Document()

    para = doc.add_paragraph()

    _add_external_hyperlink(para, "LinkedIn", "https://www.linkedin.com/in/jane-doe")



    buf = BytesIO()

    doc.save(buf)

    result = _extract_docx(buf.getvalue())



    assert "LinkedIn [https://www.linkedin.com/in/jane-doe]" in result





def test_pdf_linkedin_injected_after_anchor_word():

    page_text = "Jane Doe\nLinkedIn\nEmail: jane@example.com"

    url = "https://www.linkedin.com/in/john-doe"

    merged = _inject_pdf_hyperlinks_into_text(page_text, [url])

    assert "LinkedIn [https://www.linkedin.com/in/john-doe]" in merged





def test_pdf_linkedin_standalone_when_no_anchor():

    page_text = "Jane Doe\nSoftware Engineer"

    url = "https://www.linkedin.com/in/john-doe"

    merged = _inject_pdf_hyperlinks_into_text(page_text, [url])

    assert merged.endswith("LinkedIn [https://www.linkedin.com/in/john-doe]")





def test_pdf_extracts_annotation_uri():

    writer = PdfWriter()

    writer.add_blank_page(width=200, height=200)

    page = writer.pages[0]

    annot = DictionaryObject(

        {

            NameObject("/Type"): NameObject("/Annot"),

            NameObject("/Subtype"): NameObject("/Link"),

            NameObject("/Rect"): ArrayObject(

                [

                    NumberObject(0),

                    NumberObject(0),

                    NumberObject(100),

                    NumberObject(20),

                ]

            ),

            NameObject("/A"): DictionaryObject(

                {

                    NameObject("/S"): NameObject("/URI"),

                    NameObject("/URI"): TextStringObject(

                        "https://www.linkedin.com/in/john-doe"

                    ),

                }

            ),

        }

    )

    page[NameObject("/Annots")] = ArrayObject([writer._add_object(annot)])



    buf = BytesIO()

    writer.write(buf)



    result = _extract_pdf(buf.getvalue())



    assert "https://www.linkedin.com/in/john-doe" in result

    assert "LinkedIn [https://www.linkedin.com/in/john-doe]" in result





def test_extract_text_from_file_docx_roundtrip():

    doc = Document()

    para = doc.add_paragraph("Jane Doe")

    _add_external_hyperlink(para, "Profile", "https://www.linkedin.com/in/jane-doe")

    buf = BytesIO()

    doc.save(buf)



    result = extract_text_from_file(buf.getvalue(), "resume.docx")

    assert "Jane Doe" in result

    assert "Profile [https://www.linkedin.com/in/jane-doe]" in result





def test_paragraph_without_hyperlink_unchanged():

    doc = Document()

    doc.add_paragraph("Plain text only")

    para = doc.paragraphs[0]

    assert _paragraph_text_with_hyperlinks(para) == "Plain text only"





def test_docx_table_before_paragraph_preserves_order():

    doc = Document()

    table = doc.add_table(rows=1, cols=1)

    table.rows[0].cells[0].text = "Jane Doe\njane@example.com"

    doc.add_paragraph("Professional experience at Acme Corp")



    buf = BytesIO()

    doc.save(buf)

    result = _extract_docx(buf.getvalue())



    assert result.index("Jane Doe") < result.index("Professional experience")





def test_docx_first_page_header_extracted():

    doc = Document()

    section = doc.sections[0]

    section.different_first_page_header_footer = True

    header = section.first_page_header

    header.paragraphs[0].text = "0521234567 | contact@example.com"

    doc.add_paragraph("Body content below header")



    buf = BytesIO()

    doc.save(buf)

    result = _extract_docx(buf.getvalue())



    assert DOCX_HEADER_START in result

    assert DOCX_HEADER_END in result

    assert "contact@example.com" in result

    assert result.index("contact@example.com") < result.index("Body content")


